"""Tests for business-document text extraction (PDF/DOCX/TXT)."""

from __future__ import annotations

import io
from types import SimpleNamespace

import docx
import pytest

from app.agents.knowledge_base.file_extraction import (
    UnsupportedFileTypeError,
    extract_text_from_upload,
)


def test_extract_text_from_txt():
    text = extract_text_from_upload("notes.txt", b"Hello, this is a plain text note.")
    assert text == "Hello, this is a plain text note."


def test_extract_text_from_md():
    text = extract_text_from_upload("notes.md", "# Heading\n\nBody text.".encode("utf-8"))
    assert "Body text." in text


def test_extract_text_from_docx_round_trip():
    document = docx.Document()
    document.add_paragraph("First paragraph of a real DOCX file.")
    document.add_paragraph("Second paragraph, still real.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text_from_upload("report.docx", buffer.getvalue())

    assert "First paragraph of a real DOCX file." in text
    assert "Second paragraph, still real." in text


def test_extract_text_from_docx_skips_empty_paragraphs():
    document = docx.Document()
    document.add_paragraph("Real content.")
    document.add_paragraph("")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text_from_upload("report.docx", buffer.getvalue())

    assert text.strip() == "Real content."


def test_extract_text_from_pdf(monkeypatch):
    import app.agents.knowledge_base.file_extraction as file_extraction_module

    fake_pages = [SimpleNamespace(extract_text=lambda: "Page one text."), SimpleNamespace(extract_text=lambda: "Page two text.")]
    monkeypatch.setattr(
        file_extraction_module, "PdfReader", lambda buf: SimpleNamespace(pages=fake_pages)
    )

    text = extract_text_from_upload("doc.pdf", b"%PDF-fake-bytes")

    assert "Page one text." in text
    assert "Page two text." in text


def test_extract_text_from_pdf_handles_pages_with_no_extractable_text(monkeypatch):
    import app.agents.knowledge_base.file_extraction as file_extraction_module

    fake_pages = [SimpleNamespace(extract_text=lambda: None)]
    monkeypatch.setattr(
        file_extraction_module, "PdfReader", lambda buf: SimpleNamespace(pages=fake_pages)
    )

    text = extract_text_from_upload("empty.pdf", b"%PDF-fake-bytes")

    assert text == ""


def test_extract_text_raises_for_unsupported_extension():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_upload("archive.zip", b"whatever")
